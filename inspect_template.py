#!/usr/bin/env python3
"""
模板结构检查工具
运行后选择 Word 模板，打印表格的真实行列结构，帮助确定正确的索引。
"""
import sys
import os

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document

import tkinter as tk
from tkinter import filedialog


def inspect_template(path):
    doc = Document(path)
    if not doc.tables:
        print("模板中没有找到表格！")
        return

    for t_idx, table in enumerate(doc.tables):
        print(f"\n{'='*60}")
        print(f"表格 {t_idx}: {len(table.rows)} 行")
        print(f"{'='*60}")

        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            print(f"\n行 {r_idx} ({len(cells)} 列):")
            for c_idx, cell in enumerate(cells):
                text = cell.text.strip().replace('\n', ' | ')
                if len(text) > 40:
                    text = text[:40] + "..."
                print(f"  [{r_idx},{c_idx}] = \"{text}\"")


def main():
    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        root = tk.Tk()
        root.withdraw()
        path = filedialog.askopenfilename(
            title="选择考勤表模板",
            filetypes=[("Word 文件", "*.docx")]
        )
    if not path:
        print("未选择文件")
        return
    print(f"检查模板: {path}")
    inspect_template(path)


if __name__ == "__main__":
    main()
