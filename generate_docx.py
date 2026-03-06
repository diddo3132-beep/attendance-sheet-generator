#!/usr/bin/env python3
"""
考勤表生成模块
根据模板真实结构填充数据（通过 inspect_template.py 确认的索引）

模板结构（27行 x 13列）：
  行0: [0-1]课程名称 [2-8]值 [9-10]课程代码 [11-12]值
  行1: [0-1]教学内容 [2-12]值
  行2: [0-1]参训人数 [2-3]值 [4-5]培训课时 [6-8]值 [9-10]培训日期 [11-12]值
  行3: [0-1]教员姓名 [2-3]值 [4-5]教员签字确认 [6-8]值 [9-10]培训时间 [11-12]值
  行4: [0-1]学员作风记录 [2-3]复选框 [4-8]签字确认(合并) [9-10]培训地点 [11-12]值
  行5: 学员作风记录填写说明（不动）
  行6: 表头：序号/工号/姓名/签到（不动）
  行7-26: 学员数据（左1-20，右21-40）
"""
import os
import sys

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document


FIELD_MAP = {
    'course_name':   (0, 2),
    'course_code':   (0, 11),
    'content':       (1, 2),
    'student_count': (2, 2),
    'hours':         (2, 6),
    'date':          (2, 11),
    'instructor':    (3, 2),
    'time':          (3, 11),
    'location':      (4, 11),
}

STUDENT_START_ROW = 7
STUDENT_ROWS = 20
LEFT_ID_COL = 1
LEFT_NAME_COL = 3
RIGHT_ID_COL = 8
RIGHT_NAME_COL = 10


def set_cell(table, row_idx, col_idx, value):
    """安全地设置单元格内容，保留原有格式"""
    try:
        cell = table.rows[row_idx].cells[col_idx]
        if cell.paragraphs and cell.paragraphs[0].runs:
            p = cell.paragraphs[0]
            p.runs[0].text = str(value)
            for run in p.runs[1:]:
                run.text = ""
        else:
            cell.text = str(value)
    except IndexError:
        print(f"  警告: 索引 [{row_idx},{col_idx}] 超出范围，跳过")


def fill_attendance_template(template_path, output_path, course_info, user_info):
    doc = Document(template_path)

    if not doc.tables:
        raise ValueError("模板中没有表格")

    table = doc.tables[0]
    total_rows = len(table.rows)
    print(f"模板表格: {total_rows} 行, {len(table.rows[0].cells)} 列")

    set_cell(table, *FIELD_MAP['course_name'], course_info.get('name', ''))
    set_cell(table, *FIELD_MAP['course_code'], course_info.get('code', ''))

    content = course_info.get('content', '')
    if len(content) > 200:
        content = content[:200] + "..."
    set_cell(table, *FIELD_MAP['content'], content)

    students = user_info.get('students', [])
    set_cell(table, *FIELD_MAP['student_count'], str(len(students)))
    set_cell(table, *FIELD_MAP['hours'], course_info.get('hours', ''))
    set_cell(table, *FIELD_MAP['date'], user_info.get('date', ''))
    set_cell(table, *FIELD_MAP['instructor'], user_info.get('instructor', ''))
    set_cell(table, *FIELD_MAP['time'], user_info.get('time', ''))
    set_cell(table, *FIELD_MAP['location'], user_info.get('location', ''))

    for i in range(STUDENT_ROWS):
        row_idx = STUDENT_START_ROW + i
        if row_idx >= total_rows:
            break

        left_idx = i
        if left_idx < len(students):
            set_cell(table, row_idx, LEFT_ID_COL, students[left_idx]['id'])
            set_cell(table, row_idx, LEFT_NAME_COL, students[left_idx]['name'])

        right_idx = i + STUDENT_ROWS
        if right_idx < len(students):
            set_cell(table, row_idx, RIGHT_ID_COL, students[right_idx]['id'])
            set_cell(table, row_idx, RIGHT_NAME_COL, students[right_idx]['name'])

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"考勤表已生成: {output_path}")
