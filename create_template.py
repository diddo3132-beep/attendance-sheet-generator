#!/usr/bin/env python3
"""
生成通用考勤表模板
运行后会在 template/ 目录下创建 考勤表模板.docx，供本工具使用。
"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, 'template', '考勤表模板.docx')


def create_template():
    doc = Document()
    doc.add_paragraph('培训考勤表', style='Heading 1')

    # 27 行 x 13 列，与 generate_docx 中的索引对应
    table = doc.add_table(rows=27, cols=13)
    table.style = 'Table Grid'

    # 行 0: 课程名称、课程代码
    table.rows[0].cells[0].text = '课程名称'
    table.rows[0].cells[1].text = '课程名称'
    table.rows[0].cells[9].text = '课程代码'
    table.rows[0].cells[10].text = '课程代码'

    # 行 1: 教学内容
    table.rows[1].cells[0].text = '教学内容'
    table.rows[1].cells[1].text = '教学内容'

    # 行 2: 参训人数、培训课时、培训日期
    table.rows[2].cells[0].text = '参训人数'
    table.rows[2].cells[1].text = '参训人数'
    table.rows[2].cells[4].text = '培训课时'
    table.rows[2].cells[5].text = '培训课时'
    table.rows[2].cells[9].text = '培训日期'
    table.rows[2].cells[10].text = '培训日期'

    # 行 3: 教员姓名、教员签字确认、培训时间
    table.rows[3].cells[0].text = '教员姓名'
    table.rows[3].cells[1].text = '教员姓名'
    table.rows[3].cells[4].text = '教员签字确认'
    table.rows[3].cells[5].text = '教员签字确认'
    table.rows[3].cells[9].text = '培训时间'
    table.rows[3].cells[10].text = '培训时间'

    # 行 4: 学员记录、培训地点
    table.rows[4].cells[0].text = '学员记录'
    table.rows[4].cells[1].text = '学员记录'
    table.rows[4].cells[9].text = '培训地点'
    table.rows[4].cells[10].text = '培训地点'

    # 行 5: 填写说明（通用）
    table.rows[5].cells[0].text = '填写说明'
    table.rows[5].cells[1].text = '请在此处填写需要学员或教员注意的事项。'

    # 行 6: 表头
    table.rows[6].cells[0].text = '序号'
    table.rows[6].cells[1].text = '工号'
    table.rows[6].cells[2].text = '工号'
    table.rows[6].cells[3].text = '姓名（正楷）'
    table.rows[6].cells[4].text = '姓名（正楷）'
    table.rows[6].cells[5].text = '签到'
    table.rows[6].cells[6].text = '签到'
    table.rows[6].cells[7].text = '序号'
    table.rows[6].cells[8].text = '工号'
    table.rows[6].cells[9].text = '工号'
    table.rows[6].cells[10].text = '姓名（正楷）'
    table.rows[6].cells[11].text = '姓名（正楷）'
    table.rows[6].cells[12].text = '签到'

    # 行 7-26: 学员行，左序号 1-20，右序号 21-40
    for i in range(20):
        row_idx = 7 + i
        table.rows[row_idx].cells[0].text = str(i + 1)
        table.rows[row_idx].cells[7].text = str(i + 21)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"已生成通用模板: {OUTPUT}")


if __name__ == '__main__':
    create_template()
